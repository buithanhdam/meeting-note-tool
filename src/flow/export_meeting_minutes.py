import os
import docx
from src.config import GlobalConfig
from src.prompts import INSTRUCTIONS_CREATE_MEETING_MINUTES, SYSTEM_PROMPT, EXAMPLE_OUTPUT
from llama_index.llms.gemini import Gemini
from llama_index.core.llms import ChatMessage
import markdown
from dotenv import load_dotenv
load_dotenv()
from src.logger import get_formatted_logger
logger = get_formatted_logger(__name__)

global_config = GlobalConfig()

def _extract_response(response):
    """Extract text from model response."""
    try:
        logger.debug("Extracting response from LLM")
        if hasattr(response, 'text'):
            return response.text
        elif hasattr(response, 'content'):
            return response.content.parts[0].text
        else:
            return response.message.content
    except Exception as e:
        logger.warning(f"Exception while extracting response: {str(e)}")
        return response.message.content

def export_meeting_minutes(transcript_path):
    """Process transcript into meeting minutes"""
    try:
        logger.info(f"Generating meeting minutes from transcript: {transcript_path}")
        # Read transcript file
        with open(transcript_path, 'r') as f:
            transcript_text = f.read()
        
        logger.info(f"Transcript loaded, length: {len(transcript_text)} characters")
        
        # Use LlamaIndex and GPT to summarize
        logger.info(f"Initializing LLM with model: {global_config.GEMINI_CONFIG.model_id}")
        llm = Gemini(
            model=os.environ.get('GOOGLE_MODEL'),
            api_key=os.environ.get('GOOGLE_API_KEY')  # Corrected from using model as API key
        )
        
        messages = [
            ChatMessage(
                role="system", content=SYSTEM_PROMPT
            ),
            ChatMessage(
                role="system", content=INSTRUCTIONS_CREATE_MEETING_MINUTES
            ),
            ChatMessage(
                role="assistant", content=EXAMPLE_OUTPUT
            ),
            ChatMessage(role="user", content="Meeting transcript text: " + transcript_text),
        ]
        
        logger.info("Calling LLM to generate meeting minutes")
        resp = llm.chat(messages)
        
        minutes = _extract_response(resp)
        logger.info(f"Meeting minutes generated, length: {len(minutes)} characters")
        
        return minutes
    except Exception as e:
        logger.error(f"Error generating meeting minutes: {str(e)}")
        raise

def export_to_word(meeting_minutes_markdown: str, output_path: str = None):
    """Export meeting minutes in Markdown format to a Word (.docx) file."""
    try:
        logger.info("Converting markdown to HTML for Word export")
        # Convert Markdown to HTML
        html_content = markdown.markdown(meeting_minutes_markdown)
        
        # Create a new Word document
        logger.info("Creating Word document")
        doc = docx.Document()
        
        # Split the converted HTML into lines and process them
        logger.debug("Processing markdown content into Word format")
        for line in html_content.split('\n'):
            line = line.strip()
            if line.startswith("<h1>"):
                doc.add_heading(line.replace("<h1>", "").replace("</h1>", ""), level=1)
            elif line.startswith("<h2>"):
                doc.add_heading(line.replace("<h2>", "").replace("</h2>", ""), level=2)
            elif line.startswith("<h3>"):
                doc.add_heading(line.replace("<h3>", "").replace("</h3>", ""), level=3)
            elif line.startswith("<ul>") or line.startswith("<ol>"):
                continue  # Ignore opening list tags
            elif line.startswith("<li>"):
                doc.add_paragraph(line.replace("<li>", "- ").replace("</li>", ""), style="ListBullet")
            else:
                # Skip empty paragraphs and HTML tags
                cleaned_line = line.replace("<p>", "").replace("</p>", "").strip()
                if cleaned_line and not (cleaned_line.startswith("<") and cleaned_line.endswith(">")):
                    doc.add_paragraph(cleaned_line)
        
        # Save the document
        if not output_path:
            output_path = os.path.join(global_config.PathConfig.output_path, "meeting_minutes.docx")
        
        # Ensure directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc.save(output_path)
        logger.info(f"Word document saved to: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Error exporting to Word: {str(e)}")
        raise